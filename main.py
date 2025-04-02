import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, font
from docx import Document
from fpdf import FPDF
import os

class TextEditor:
    def __init__(self, root):
        self.root = root
        self.root.title("Текстовый редактор")
        self.text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, font=("Arial", 12))
        self.text_area.pack(expand=True, fill="both")

        self.create_menu()
        
    def create_menu(self):
        menu_bar = tk.Menu(self.root)
        
        file_menu = tk.Menu(menu_bar, tearoff=0)
        file_menu.add_command(label="Открыть", command=self.open_file)
        file_menu.add_command(label="Сохранить", command=self.save_file)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self.root.quit)
        menu_bar.add_cascade(label="Файл", menu=file_menu)

        format_menu = tk.Menu(menu_bar, tearoff=0)
        format_menu.add_command(label="Жирный", command=self.bold_text)
        format_menu.add_command(label="Курсив", command=self.italic_text)
        format_menu.add_command(label="Подчеркнутый", command=self.underline_text)
        menu_bar.add_cascade(label="Формат", menu=format_menu)

        self.root.config(menu=menu_bar)

    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Текстовые файлы", "*.txt;*.docx;*.rtf")])
        if not file_path:
            return
        self.text_area.delete(1.0, tk.END)

        if file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as file:
                self.text_area.insert(tk.END, file.read())
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            self.text_area.insert(tk.END, text)
        elif file_path.endswith(".rtf"):
            with open(file_path, "r", encoding="utf-8") as file:
                self.text_area.insert(tk.END, file.read())

    def save_file(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Текстовые файлы", "*.txt"),
                                                            ("Документы Word", "*.docx"),
                                                            ("Файлы RTF", "*.rtf"),
                                                            ("Файлы PDF", "*.pdf")])
        if not file_path:
            return
        
        content = self.text_area.get(1.0, tk.END).strip()

        if file_path.endswith(".txt"):
            with open(file_path, "w", encoding="utf-8") as file:
                file.write(content)
        elif file_path.endswith(".docx"):
            doc = Document()
            doc.add_paragraph(content)
            doc.save(file_path)
        elif file_path.endswith(".rtf"):
            with open(file_path, "w", encoding="utf-8") as file:
                file.write(content)
        elif file_path.endswith(".pdf"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(190, 10, content)
            pdf.output(file_path)

    def bold_text(self):
        self.apply_font_style(weight="bold")

    def italic_text(self):
        self.apply_font_style(slant="italic")

    def underline_text(self):
        self.apply_font_style(underline=True)

    def apply_font_style(self, weight="normal", slant="roman", underline=False):
        current_font = font.Font(self.text_area, self.text_area.cget("font"))
        new_font = font.Font(family=current_font.actual()["family"],
                             size=current_font.actual()["size"],
                             weight=weight, slant=slant, underline=underline)
        self.text_area.config(font=new_font)

if __name__ == "__main__":
    root = tk.Tk()
    app = TextEditor(root)
    root.mainloop()
